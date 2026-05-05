"""
Build DC (Descrição de Cargo) template from the prototype docx.
Replaces metadata fields with eigenpal placeholders and clears content sections.
"""
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:\Users\leandro.theodoro.MN-NTB-LEANDROT\Downloads\DC_Prototipo_A_Editorial.docx"
OUT = r"C:\Users\leandro.theodoro.MN-NTB-LEANDROT\Downloads\DC_Template_Descricao_Cargo.docx"

doc = Document(SRC)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def para_text(p):
    return "".join(r.text for r in p.runs)

def set_para_text(p, new_text):
    """Replace paragraph text, preserving formatting of the first non-empty run."""
    runs = p.runs
    if not runs:
        return
    # Find first run that has content (for formatting reference)
    ref_run = next((r for r in runs if r.text.strip()), runs[0])
    ref_run.text = new_text
    for r in runs:
        if r is not ref_run:
            r.text = ""

def replace_para(p, replacements):
    full = para_text(p)
    result = full
    for old, new in replacements.items():
        result = result.replace(old, new)
    if result != full:
        set_para_text(p, result)

def clear_para(p):
    set_para_text(p, "")

def replace_in_section_parts(section_obj, replacements):
    """Apply replacements to header/footer paragraphs and their tables."""
    for part in [section_obj.header, section_obj.footer,
                 section_obj.first_page_header, section_obj.first_page_footer]:
        if part is None:
            continue
        for p in part.paragraphs:
            replace_para(p, replacements)
        for t in part.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_para(p, replacements)

def clear_table_value_col(table, value_col=1):
    """Clear the value column in a key-value table (skip col 0 which is the label)."""
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci == value_col:
                for p in cell.paragraphs:
                    clear_para(p)

def clear_table_data_rows(table, header_rows=1):
    """Clear all rows after the header rows."""
    for ri, row in enumerate(table.rows):
        if ri < header_rows:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                clear_para(p)

def replace_table_value_col(table, replacements, value_col=1):
    """Apply replacements only to the value column of a key-value table."""
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci == value_col:
                for p in cell.paragraphs:
                    replace_para(p, replacements)

# ---------------------------------------------------------------------------
# 1. Cover: title paragraph and subtitle
# ---------------------------------------------------------------------------
# Paragraph [3]: "Operador de Caixa"  →  "{doc_title}"
# Paragraph [4]: "DC-01  ·  Revisão 00  ·  Área Financeiro"  →  compound placeholder
CONTENT_PARAS = doc.paragraphs

set_para_text(CONTENT_PARAS[3], "{doc_title}")
set_para_text(CONTENT_PARAS[4], "{doc_code}  ·  Revisão {revision_number}  ·  Área {controlled_by_area}")

# ---------------------------------------------------------------------------
# 2. Table 0 — Identification table (rows: CÓDIGO REVISÃO TÍTULO ÁREA EMISSÃO ÚLT.REV ELAB APROV)
# ---------------------------------------------------------------------------
id_table = doc.tables[0]
# Row index → placeholder for value cell (col 1)
id_replacements = {
    0: "{doc_code}",
    1: "{revision_number}",
    2: "{doc_title}",
    3: "{controlled_by_area}",
    4: "{effective_date}",
    5: "{effective_date}",
    6: "{author}",
    7: "{approvers}",
}
for row_idx, placeholder in id_replacements.items():
    cell = id_table.rows[row_idx].cells[1]
    for p in cell.paragraphs:
        set_para_text(p, placeholder)

# ---------------------------------------------------------------------------
# 3. Section 01 — Mission paragraph (paragraph [7])
# ---------------------------------------------------------------------------
clear_para(CONTENT_PARAS[7])

# ---------------------------------------------------------------------------
# 4. Table 1 — Posição na Estrutura: clear value column
# ---------------------------------------------------------------------------
clear_table_value_col(doc.tables[1], value_col=1)

# ---------------------------------------------------------------------------
# 5. Section 03 — Escopo paragraphs: keep bold label, clear content text
#    Para [10]: "O que o cargo cobre.  <content>"
#    Para [11]: "O que o cargo não cobre.  <content>"
# ---------------------------------------------------------------------------
def trim_para_to_bold_label(p):
    """Keep only bold runs, clear the rest."""
    for r in p.runs:
        if not r.bold:
            r.text = ""

trim_para_to_bold_label(CONTENT_PARAS[10])
trim_para_to_bold_label(CONTENT_PARAS[11])

# ---------------------------------------------------------------------------
# 6. Table 2 — Responsabilidades: clear data rows (keep header row 0)
# ---------------------------------------------------------------------------
clear_table_data_rows(doc.tables[2], header_rows=1)

# ---------------------------------------------------------------------------
# 7. Table 3 — Autoridades: clear data rows
# ---------------------------------------------------------------------------
clear_table_data_rows(doc.tables[3], header_rows=1)

# ---------------------------------------------------------------------------
# 8. Table 4 — Interfaces: clear data rows
# ---------------------------------------------------------------------------
clear_table_data_rows(doc.tables[4], header_rows=1)

# ---------------------------------------------------------------------------
# 9. Table 5 — Indicadores: clear data rows
# ---------------------------------------------------------------------------
clear_table_data_rows(doc.tables[5], header_rows=1)

# ---------------------------------------------------------------------------
# 10. Table 6 — Requisitos do Cargo: clear value column (col 1)
# ---------------------------------------------------------------------------
clear_table_value_col(doc.tables[6], value_col=1)

# ---------------------------------------------------------------------------
# 11. Table 7 — Competências: clear data rows
# ---------------------------------------------------------------------------
clear_table_data_rows(doc.tables[7], header_rows=1)

# ---------------------------------------------------------------------------
# 12. Table 8 — Diretrizes: clear data rows
# ---------------------------------------------------------------------------
clear_table_data_rows(doc.tables[8], header_rows=1)

# ---------------------------------------------------------------------------
# 13. Table 9 — Histórico de Revisões: keep header + seed first revision row
# ---------------------------------------------------------------------------
hist_table = doc.tables[9]
if len(hist_table.rows) >= 2:
    hist_row = hist_table.rows[1]
    # Col 0: REVISÃO = {revision_number}
    set_para_text(hist_row.cells[0].paragraphs[0], "{revision_number}")
    # Col 1: DATA = {effective_date}
    set_para_text(hist_row.cells[1].paragraphs[0], "{effective_date}")
    # Col 2: DESCRIÇÃO keep "Emissão inicial do documento."
    # Col 3: RESPONSÁVEL = {author}
    set_para_text(hist_row.cells[3].paragraphs[0], "{author}")

# ---------------------------------------------------------------------------
# 14. Headers and footers — replace metadata values with placeholders
# ---------------------------------------------------------------------------
HEADER_FOOTER_REPLACEMENTS = {
    "DC-01": "{doc_code}",
    "Operador de Caixa": "{doc_title}",
    "[Definir internamente]": "{author}",
    "Gerente Financeiro": "{approvers}",
    "__/__/____": "{effective_date}",
    # Revision number as it appears in the running header cell
    "REVISÃO  00": "REVISÃO  {revision_number}",
    "REVISÃO 00": "REVISÃO {revision_number}",
}

for section in doc.sections:
    replace_in_section_parts(section, HEADER_FOOTER_REPLACEMENTS)

# After "Gerente Financeiro" is replaced everywhere, "Financeiro" alone only
# remains in the ÁREA context. Replace it in headers/footers now.
AREA_REPLACEMENT = {"Financeiro": "{controlled_by_area}"}
for section in doc.sections:
    replace_in_section_parts(section, AREA_REPLACEMENT)

# Handle revision "00" in headers — it appears as a standalone value next to REVISÃO label.
# Replace only the exact string " 00" preceded by "REVISÃO" context.
# Simpler: do a targeted replace of "00" only where it's alone in a paragraph.
def replace_standalone_revision(section_obj):
    for part in [section_obj.header, section_obj.footer,
                 section_obj.first_page_header, section_obj.first_page_footer]:
        if part is None:
            continue
        for t in part.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        txt = para_text(p).strip()
                        if txt == "00":
                            set_para_text(p, "{revision_number}")

for section in doc.sections:
    replace_standalone_revision(section)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT)
print(f"Template saved: {OUT}")
